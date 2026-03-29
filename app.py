import streamlit as st
import requests
import json
import io
import re
import time
import pandas as pd
from datetime import datetime
from supabase import create_client

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.oxml import parse_xml
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

try:
    from pypdf import PdfReader
    PDF_OK = True
except ImportError:
    PDF_OK = False

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="JobFit AI",
    page_icon="🎯",
    layout="wide",
    initial_sidebar_state="collapsed",
)

# ── Google Auth ──────────────────────────────────────────────────────────────
if not st.user.is_logged_in:
    st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Serif+Display&display=swap');
html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; }
.stApp, .main, [data-testid="stAppViewContainer"] { background: #f0f2f8 !important; }
[data-testid="stHeader"] { background: #ffffff !important; border-bottom: 1px solid #dde1ef !important; }
.block-container { padding: 2rem 2.5rem 4rem; max-width: 1100px; }
.lp-hero {
    background: #ffffff;
    border: 1px solid #dde1ef;
    border-radius: 18px;
    padding: 3.5rem 2rem 2.8rem;
    text-align: center;
    margin-bottom: 1.5rem;
    box-shadow: 0 2px 8px rgba(37,99,235,0.06);
    position: relative;
    overflow: hidden;
}
.lp-hero::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 4px;
    background: linear-gradient(90deg, #2563eb, #7c3aed, #0891b2);
}
.lp-hero h1 { font-family: 'DM Serif Display', serif; font-size: 3rem; color: #1a1d2e; margin: 0; }
.lp-hero h1 span { color: #2563eb; }
.lp-hero .tagline { color: #64748b; font-size: 1.15rem; margin-top: 0.75rem; margin-bottom: 0.5rem; }
.lp-hero .sub { color: #94a3b8; font-size: 0.95rem; margin-bottom: 2rem; }
.lp-features {
    display: grid;
    grid-template-columns: repeat(3, 1fr);
    gap: 1rem;
    margin-bottom: 1.5rem;
}
.lp-feature {
    background: #ffffff;
    border: 1px solid #dde1ef;
    border-radius: 14px;
    padding: 1.5rem;
    text-align: center;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.lp-feature .icon { font-size: 2rem; margin-bottom: 0.75rem; }
.lp-feature h3 { font-size: 1rem; font-weight: 600; color: #1a1d2e; margin: 0 0 0.4rem; }
.lp-feature p { font-size: 0.85rem; color: #64748b; margin: 0; }
.lp-steps {
    background: #ffffff;
    border: 1px solid #dde1ef;
    border-radius: 14px;
    padding: 1.6rem 2rem;
    margin-bottom: 1.5rem;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.lp-steps-title { font-weight: 600; font-size: 0.85rem; letter-spacing: 0.08em; text-transform: uppercase; color: #64748b; margin-bottom: 1rem; }
.lp-steps-row { display: flex; align-items: center; justify-content: center; gap: 0.5rem; flex-wrap: wrap; }
.lp-step { text-align: center; min-width: 100px; }
.lp-step .s-icon { font-size: 1.5rem; margin-bottom: 0.3rem; }
.lp-step .s-title { font-size: 0.8rem; font-weight: 600; color: #1a1d2e; }
.lp-step .s-desc { font-size: 0.72rem; color: #94a3b8; }
.lp-arrow { color: #dde1ef; font-size: 1.2rem; }
.lp-signin {
    background: #ffffff;
    border: 1px solid #dde1ef;
    border-radius: 14px;
    padding: 2rem;
    text-align: center;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.lp-signin p { color: #64748b; font-size: 0.9rem; margin-bottom: 1rem; }
.lp-signin .note { color: #94a3b8; font-size: 0.78rem; margin-top: 0.75rem; }
</style>
""", unsafe_allow_html=True)

    st.markdown("""
<div class="lp-hero">
    <h1>🎯 <span>JobFit AI</span></h1>
    <p class="tagline">Know your score. Apply smart.</p>
    <p class="sub">Upload your resume &amp; job description — get an instant match score, gap analysis, and tailored cover letter.</p>
</div>
""", unsafe_allow_html=True)

    st.markdown("""
<div class="lp-features">
    <div class="lp-feature">
        <div class="icon">🎯</div>
        <h3>Match Score</h3>
        <p>Get a 0–100% score showing how well your resume matches the job description.</p>
    </div>
    <div class="lp-feature">
        <div class="icon">🔍</div>
        <h3>Gap Analysis</h3>
        <p>See exactly which skills and experience you have, partially have, or are missing.</p>
    </div>
    <div class="lp-feature">
        <div class="icon">✉️</div>
        <h3>Cover Letter</h3>
        <p>Get a tailored cover letter generated instantly — ready to send or customize.</p>
    </div>
</div>
""", unsafe_allow_html=True)

    st.markdown("""
<div class="lp-steps">
    <div class="lp-steps-title">How it works</div>
    <div class="lp-steps-row">
        <div class="lp-step">
            <div class="s-icon">📄</div>
            <div class="s-title">Upload Resume</div>
            <div class="s-desc">PDF, DOCX or TXT</div>
        </div>
        <div class="lp-arrow">→</div>
        <div class="lp-step">
            <div class="s-icon">💼</div>
            <div class="s-title">Add Job Description</div>
            <div class="s-desc">Upload or paste it</div>
        </div>
        <div class="lp-arrow">→</div>
        <div class="lp-step">
            <div class="s-icon">🤖</div>
            <div class="s-title">AI Analyzes</div>
            <div class="s-desc">Compares both instantly</div>
        </div>
        <div class="lp-arrow">→</div>
        <div class="lp-step">
            <div class="s-icon">📊</div>
            <div class="s-title">Get Your Score</div>
            <div class="s-desc">0–100% match rating</div>
        </div>
        <div class="lp-arrow">→</div>
        <div class="lp-step">
            <div class="s-icon">⬇️</div>
            <div class="s-title">Download Files</div>
            <div class="s-desc">Report, cover letter &amp; resume</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

    st.markdown("""
<div class="lp-signin">
    <p>Sign in with your Google account to get started — free, no credit card needed.</p>
</div>
""", unsafe_allow_html=True)

    col1, col2, col3 = st.columns([1.5, 1, 1.5])
    with col2:
        if st.button("🔐 Sign in with Google", use_container_width=True):
            st.login("google")

    st.markdown("""
<div style="text-align:center; color:#94a3b8; font-size:0.78rem; margin-top:0.75rem;">
    🔒 Secure login via Google &nbsp;·&nbsp; Your data stays private &nbsp;·&nbsp; No spam ever
</div>
""", unsafe_allow_html=True)

    st.stop()

# ── Logged in — get user info ─────────────────────────────────────────────────
user_email = st.user.email or ""
user_name  = st.user.name or "User"

# ── Supabase client (shared across all tabs) ──────────────────────────────────
supabase = create_client(
    st.secrets["SUPABASE_URL"],
    st.secrets["SUPABASE_KEY"]
)

# ── CSS ───────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Sans:wght@300;400;500;600;700&family=DM+Serif+Display&display=swap');

:root {
    --bg:         #f0f2f8;
    --bg2:        #e8ebf5;
    --white:      #ffffff;
    --border:     #dde1ef;
    --text:       #1a1d2e;
    --muted:      #64748b;
    --primary:    #2563eb;
    --primary-lt: #eff6ff;
    --green:      #059669;
    --amber:      #d97706;
    --red:        #dc2626;
}

html, body, [class*="css"] { font-family: 'DM Sans', sans-serif; color: var(--text); }
.stApp, .main, [data-testid="stAppViewContainer"] { background: var(--bg) !important; }
[data-testid="stHeader"] { background: var(--white) !important; border-bottom: 1px solid var(--border) !important; }
.block-container { padding: 2rem 2.5rem 4rem; max-width: 1280px; }

.hero {
    background: var(--white);
    border: 1px solid var(--border);
    border-radius: 18px;
    padding: 2.8rem 2rem 2.2rem;
    text-align: center;
    margin-bottom: 1.5rem;
    box-shadow: 0 2px 8px rgba(37,99,235,0.06);
    position: relative;
    overflow: hidden;
}
.hero::before {
    content: '';
    position: absolute;
    top: 0; left: 0; right: 0; height: 4px;
    background: linear-gradient(90deg, #2563eb, #7c3aed, #0891b2);
}
.hero h1 { font-family: 'DM Serif Display', serif; font-size: 2.8rem; color: var(--text); margin: 0; }
.hero h1 span { color: var(--primary); }
.hero p { color: var(--muted); margin-top: 0.5rem; font-size: 1rem; }

.hiw-wrap {
    background: var(--white);
    border: 1px solid var(--border);
    border-radius: 14px;
    padding: 1.6rem 2rem;
    margin-bottom: 1.8rem;
    box-shadow: 0 1px 4px rgba(0,0,0,0.04);
}
.hiw-title {
    font-weight: 700; font-size: 0.85rem; text-transform: uppercase;
    letter-spacing: 0.08em; color: var(--muted); margin-bottom: 1rem;
}
.hiw-steps { display: flex; gap: 0; align-items: stretch; }
.hiw-step {
    flex: 1; display: flex; flex-direction: column; align-items: center;
    text-align: center; padding: 0 1rem; position: relative;
}
.hiw-step:not(:last-child)::after {
    content: '→'; position: absolute; right: -10px; top: 14px;
    font-size: 1.2rem; color: #cbd5e1;
}
.hiw-icon {
    width: 40px; height: 40px; background: var(--primary-lt);
    border: 2px solid #bfdbfe; border-radius: 50%;
    display: flex; align-items: center; justify-content: center;
    font-size: 1rem; margin-bottom: 0.5rem;
}
.hiw-step-title { font-weight: 700; font-size: 0.85rem; color: var(--text); }
.hiw-step-desc  { font-size: 0.75rem; color: var(--muted); margin-top: 2px; line-height: 1.4; }

.demo-banner {
    background: linear-gradient(135deg, #eff6ff, #f5f3ff);
    border: 1.5px solid #bfdbfe;
    border-radius: 12px;
    padding: 1rem 1.4rem;
    margin-bottom: 1.2rem;
    display: flex; align-items: center; gap: 1rem;
}
.demo-banner-text strong { color: var(--primary); font-size: 0.95rem; display: block; }
.demo-banner-text p { color: var(--muted); font-size: 0.8rem; margin: 2px 0 0; }

.card {
    background: var(--white); border: 1px solid var(--border); border-radius: 12px;
    padding: 1.4rem 1.6rem; box-shadow: 0 1px 4px rgba(0,0,0,0.04); margin-bottom: 1rem;
}
.card h4 {
    font-size: 0.72rem; font-weight: 700; text-transform: uppercase;
    letter-spacing: 0.08em; color: var(--primary); margin: 0 0 0.5rem;
}
.card p { color: #374151; margin: 0; font-size: 0.92rem; line-height: 1.65; }

.score-wrap { border-radius: 14px; padding: 2.2rem 1.5rem; text-align: center; margin-bottom: 1rem; }
.score-green { background: #d1fae5; border: 2px solid #059669; }
.score-amber { background: #fef3c7; border: 2px solid #d97706; }
.score-red   { background: #fee2e2; border: 2px solid #dc2626; }
.score-number { font-family: 'DM Serif Display', serif; font-size: 4.5rem; line-height: 1; }
.score-green .score-number { color: #065f46; }
.score-amber .score-number { color: #78350f; }
.score-red   .score-number { color: #7f1d1d; }
.score-label { font-size: 1.2rem; font-weight: 700; margin-top: 0.4rem; }
.score-green .score-label { color: #065f46; }
.score-amber .score-label { color: #78350f; }
.score-red   .score-label { color: #7f1d1d; }
.score-advice { font-size: 0.88rem; font-weight: 500; margin-top: 0.6rem; line-height: 1.5; padding: 0 0.5rem; }
.score-green .score-advice { color: #064e3b; }
.score-amber .score-advice { color: #713f12; }
.score-red   .score-advice { color: #7f1d1d; }

.metrics { display: flex; gap: 0.75rem; margin-bottom: 1.25rem; flex-wrap: wrap; }
.metric {
    flex: 1; min-width: 120px; background: var(--white); border: 1px solid var(--border);
    border-radius: 10px; padding: 0.9rem 1rem; text-align: center;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
}
.metric-val { font-size: 1.7rem; font-weight: 700; line-height: 1; }
.metric-lbl { font-size: 0.7rem; font-weight: 600; text-transform: uppercase; letter-spacing: 0.07em; color: var(--muted); margin-top: 3px; }

.pill { display: inline-block; padding: 4px 11px; border-radius: 20px; font-size: 0.78rem; font-weight: 600; margin: 3px; }
.pill-green { background: #d1fae5; color: #065f46; border: 1px solid #6ee7b7; }
.pill-amber { background: #fef3c7; color: #78350f; border: 1px solid #fcd34d; }
.pill-red   { background: #fee2e2; color: #7f1d1d; border: 1px solid #fca5a5; }

.gap-block { background: var(--white); border-radius: 10px; border-left: 4px solid; padding: 1.2rem 1.4rem; margin-bottom: 0.8rem; box-shadow: 0 1px 3px rgba(0,0,0,0.04); }
.gap-green { border-color: #059669; }
.gap-amber { border-color: #d97706; }
.gap-red   { border-color: #dc2626; }
.gap-title { font-weight: 700; font-size: 0.95rem; margin-bottom: 0.5rem; color: var(--text); }
.gap-sub   { font-size: 0.85rem; color: var(--muted); margin: 0; }

.strength {
    background: var(--primary-lt); border: 1px solid #bfdbfe; border-radius: 10px;
    padding: 1rem 1.2rem; font-size: 0.9rem; color: #1e3a8a; font-weight: 500; margin-bottom: 0.5rem;
}

.cover-box {
    background: var(--white); border: 1px solid var(--border); border-radius: 12px;
    padding: 2rem; font-size: 0.93rem; line-height: 1.85; color: var(--text);
    white-space: pre-wrap; box-shadow: 0 1px 4px rgba(0,0,0,0.05);
}

.badge {
    display: inline-flex; align-items: center; gap: 6px;
    background: #eff6ff; border: 1px solid #bfdbfe; border-radius: 8px;
    padding: 6px 14px; font-size: 0.83rem; color: #1d4ed8; font-weight: 600;
}

.dl-card { background: var(--white); border: 1px solid var(--border); border-radius: 12px; padding: 1.2rem 1.4rem; box-shadow: 0 1px 4px rgba(0,0,0,0.05); text-align: center; }
.dl-icon  { font-size: 2rem; margin-bottom: 0.4rem; }
.dl-title { font-weight: 700; font-size: 0.95rem; color: var(--text); }
.dl-desc  { font-size: 0.78rem; color: var(--muted); margin-top: 2px; }

/* App row card for My Applications tab */
.app-row {
    background: var(--white); border: 1px solid var(--border); border-radius: 12px;
    padding: 1rem 1.4rem; margin-bottom: 0.75rem;
    box-shadow: 0 1px 3px rgba(0,0,0,0.04);
    display: flex; align-items: center; gap: 1rem; flex-wrap: wrap;
}
.app-row-company { font-weight: 700; font-size: 1rem; color: var(--text); flex: 1; min-width: 140px; }
.app-row-meta { font-size: 0.78rem; color: var(--muted); }
.app-score-badge {
    display: inline-block; padding: 3px 10px; border-radius: 20px;
    font-size: 0.82rem; font-weight: 700;
}
.app-score-green { background: #d1fae5; color: #065f46; }
.app-score-amber { background: #fef3c7; color: #78350f; }
.app-score-red   { background: #fee2e2; color: #7f1d1d; }

/* Status badge */
.status-badge {
    display: inline-block; padding: 3px 10px; border-radius: 20px;
    font-size: 0.78rem; font-weight: 600;
}
.status-applied    { background: #eff6ff; color: #1d4ed8; border: 1px solid #bfdbfe; }
.status-interview  { background: #f5f3ff; color: #5b21b6; border: 1px solid #ddd6fe; }
.status-rejected   { background: #fee2e2; color: #7f1d1d; border: 1px solid #fca5a5; }
.status-offer      { background: #d1fae5; color: #065f46; border: 1px solid #6ee7b7; }

/* Nav tabs */
.stTabs [data-baseweb="tab-list"] { background: var(--white); border-radius: 10px; padding: 4px; gap: 4px; border: 1px solid var(--border); }
.stTabs [data-baseweb="tab"]      { border-radius: 7px; color: var(--muted); font-weight: 600; font-size: 0.88rem; padding: 8px 20px; }
.stTabs [aria-selected="true"]    { background: var(--primary) !important; color: #fff !important; }

[data-testid="stFileUploader"] { background: var(--white) !important; border: 2px dashed var(--border) !important; border-radius: 12px !important; }
.stTextArea textarea { background: var(--white) !important; border: 1px solid var(--border) !important; border-radius: 10px !important; color: var(--text) !important; font-size: 0.9rem !important; }
.stTextArea textarea:focus { border-color: var(--primary) !important; box-shadow: 0 0 0 3px rgba(37,99,235,0.1) !important; }
.stButton > button { border-radius: 9px !important; font-weight: 600 !important; font-size: 0.9rem !important; }
[data-testid="stDownloadButton"] button {
    width: 100% !important; border-radius: 9px !important; font-weight: 600 !important;
    background: var(--primary) !important; color: white !important; border: none !important; margin-top: 0.6rem !important;
}

hr { border-color: var(--border) !important; margin: 1.5rem 0 !important; }

/* Demo button flush against card */
div[data-testid="stButton"]:has(button[kind="primary"]#demo_top) {
    margin-top: 0 !important;
    padding-top: 0 !important;
}
</style>
""", unsafe_allow_html=True)

# ── Demo data ─────────────────────────────────────────────────────────────────
DEMO_RESUME = """John Smith
Email: john.smith@email.com | Phone: (555) 123-4567

SUMMARY
Data Analyst with 4 years of experience in SQL, Python, and Power BI. Skilled in building
dashboards, data pipelines, and delivering insights to business stakeholders.

EXPERIENCE
Senior Data Analyst — Acme Corp (2021–Present)
- Built 15+ Power BI dashboards used by 200+ employees
- Wrote complex SQL queries to analyze 10M+ row datasets
- Automated reporting with Python (pandas, matplotlib), saving 8 hrs/week

Data Analyst — StartupXYZ (2020–2021)
- Analyzed user behavior data using Python and Excel
- Created KPI reports for executive team

SKILLS
Python, SQL, Power BI, Excel, Tableau, pandas, numpy, data visualization, ETL pipelines

EDUCATION
B.S. Computer Science — State University, 2019
"""

DEMO_JD = """Company: DataDriven Inc.
Job Title: Data Analyst

We are looking for a Data Analyst to join our growing analytics team.

Requirements:
- 3+ years of experience in data analysis
- Proficiency in SQL and Python
- Experience with Power BI or Tableau
- Strong communication skills
- Experience with ETL pipelines
- Familiarity with cloud platforms (AWS/Azure) preferred
- Bachelor's degree in relevant field

Responsibilities:
- Build and maintain dashboards and reports
- Collaborate with business teams to define KPIs
- Analyze large datasets and present findings
- Support data pipeline development
"""

# ── Session state ─────────────────────────────────────────────────────────────
defaults = {
    "analysis_done": False, "score": 0, "gap_data": None,
    "cover_letter": "", "resume_text": "", "jd_text": "",
    "company_name": "", "applicant_name": "", "is_demo": False,
}
for k, v in defaults.items():
    if k not in st.session_state:
        st.session_state[k] = v


# ── File extraction ───────────────────────────────────────────────────────────
def extract_text(uploaded_file) -> str:
    name = uploaded_file.name.lower()
    if name.endswith(".txt"):
        return uploaded_file.read().decode("utf-8", errors="ignore")
    elif name.endswith(".pdf"):
        if not PDF_OK:
            st.error("Install pypdf: pip install pypdf")
            return ""
        reader = PdfReader(io.BytesIO(uploaded_file.read()))
        return "\n".join(p.extract_text() or "" for p in reader.pages)
    elif name.endswith(".docx"):
        if not DOCX_OK:
            st.error("Install python-docx: pip install python-docx")
            return ""
        doc = Document(io.BytesIO(uploaded_file.read()))
        return "\n".join(p.text for p in doc.paragraphs)
    return ""


# ── Gemini analysis ──────────────────────────────────────────────────────────
def analyze(resume: str, jd: str) -> dict:
    api_key = st.secrets["GEMINI_API_KEY"]
    url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-flash-lite:generateContent?key={api_key}"

    prompt = f"""You are an expert ATS system and career coach.
Analyze the Resume vs Job Description and return ONLY valid JSON — no markdown, no explanation, no code fences.

Scoring rules — be strict and realistic:
- 90-100%: Candidate meets every requirement perfectly, including nice-to-haves
- 80-89%: Strong match, meets all must-haves, missing only minor nice-to-haves
- 60-79%: Decent match but has notable gaps in required skills or experience
- 40-59%: Significant gaps, missing several key requirements
- Below 40%: Poor match, missing most requirements
A score above 90% should be rare. Always penalize missing required skills, lack of experience, or missing certifications.

RESUME:
{resume}

JOB DESCRIPTION:
{jd}

Return exactly this JSON:
{{
  "applicant_name": "<full name from resume or 'Applicant'>",
  "company_name": "<company name from JD or 'the Company'>",
  "score": <integer 0-100>,
  "score_reasoning": "<2-3 sentences>",
  "matched_skills": ["skill1", "skill2"],
  "partial_skills": [{{"skill":"name","resume_level":"what candidate has","required_level":"what JD needs"}}],
  "missing_skills": ["skill1"],
  "matched_experience": ["point1"],
  "missing_experience": ["gap1"],
  "education_match": "<assessment>",
  "strengths": ["strength1","strength2","strength3"],
  "improvement_suggestions": ["suggestion1","suggestion2","suggestion3"],
  "cover_letter": "<full professional cover letter 3-4 paragraphs, no placeholder brackets, signed with applicant name>"
}}"""
    payload = {"contents": [{"parts": [{"text": prompt}]}]}
    for attempt in range(3):
        resp = requests.post(url, json=payload)
        if resp.status_code == 429:
            time.sleep(10)
            continue
        if not resp.ok:
            raise Exception(f"Gemini API error {resp.status_code}: {resp.text}")
        break
    else:
        raise Exception("Gemini rate limit reached. Please wait a minute and try again.")
    raw = resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw)
    raw = re.sub(r"\s*```$", "", raw)
    return json.loads(raw)


# ── Score meta ────────────────────────────────────────────────────────────────
def score_meta(score):
    if score >= 80:
        return "score-green", "🚀 Apply Immediately!", \
               "Strong match — your profile aligns well. Submit your application right away!"
    elif score >= 60:
        return "score-amber", "🤔 Consider & Apply", \
               "Decent match but some gaps exist. Review the suggestions before applying."
    else:
        return "score-red", "⚠️ Significant Gaps", \
               "Key requirements are missing. Upskill first, or apply knowing you will need to learn fast."


# ── DOCX helpers ──────────────────────────────────────────────────────────────
def _set_cell_bg(cell, hex_color):
    from docx.oxml.ns import qn
    shading = parse_xml(f'<w:shd {cell._element.nsmap["w"]} w:val="clear" w:color="auto" w:fill="{hex_color}"/>')
    cell._element.tcPr.append(shading)

def _add_divider(doc, color="2563EB"):
    from docx.oxml.ns import qn
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="6" w:space="1" w:color="{color}"/></w:pBdr>')
    pPr.append(pBdr)
    return p

def _doc_header(doc, title, subtitle, applicant, company, date_str):
    accent = doc.add_paragraph()
    accent.paragraph_format.space_after = Pt(0)
    from docx.oxml.ns import qn
    pPr = accent._element.get_or_add_pPr()
    pBdr = parse_xml('<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:top w:val="single" w:sz="24" w:space="1" w:color="2563EB"/></w:pBdr>')
    pPr.append(pBdr)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("🎯  " + title)
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x2E)

    if subtitle:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        s.paragraph_format.space_after = Pt(4)
        sr = s.add_run(subtitle)
        sr.font.name = "Calibri"
        sr.font.size = Pt(11)
        sr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.paragraph_format.space_after = Pt(2)
    mr = meta.add_run(f"👤 {applicant}   |   🏢 {company}   |   📅 {date_str}")
    mr.font.name = "Calibri"
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    _add_divider(doc)

def _section_heading(doc, text, color_rgb=(37, 99, 235)):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(*color_rgb)
    return p

def _body_text(doc, text, italic=False, color_rgb=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.font.italic = italic
    if color_rgb:
        run.font.color.rgb = RGBColor(*color_rgb)
    return p

def _skills_table(doc, items, bg_color, text_color_rgb):
    from docx.oxml.ns import qn
    if not items:
        return
    cols = 3
    rows_needed = (len(items) + cols - 1) // cols
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.style = "Table Grid"
    for r_idx in range(rows_needed):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            item_idx = r_idx * cols + c_idx
            cell.width = Inches(2.17)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{bg_color}"/>')
            tcPr.append(shd)
            if item_idx < len(items):
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
                run = p.add_run(items[item_idx])
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(*text_color_rgb)
            else:
                cell.paragraphs[0].add_run("")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def _add_footer(doc, left_text, right_text):
    from docx.oxml.ns import qn
    section = doc.sections[0]
    footer = section.footer
    ft = footer.paragraphs[0]
    ft.clear()
    ft.paragraph_format.space_before = Pt(0)
    lr = ft.add_run(left_text)
    lr.font.size = Pt(9)
    lr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    lr.font.name = "Calibri"
    ft.add_run("\t")
    rr = ft.add_run(right_text)
    rr.font.size = Pt(9)
    rr.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    rr.font.name = "Calibri"
    pPr = ft._element.get_or_add_pPr()
    tabs = parse_xml('<w:tabs xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:tab w:val="right" w:pos="9360"/></w:tabs>')
    pPr.append(tabs)


# ── DOCX builders ─────────────────────────────────────────────────────────────
def build_report_docx(score, gap, company):
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    applicant = gap.get("applicant_name", "Applicant")
    date_str  = datetime.now().strftime("%B %d, %Y")

    _doc_header(doc, "JobFit AI — Match Report", "AI-Powered Resume Analysis", applicant, company, date_str)

    _, label, advice = score_meta(score)
    if score >= 80:
        score_bg, score_fg = "D1FAE5", (6, 95, 70)
    elif score >= 60:
        score_bg, score_fg = "FEF3C7", (120, 53, 15)
    else:
        score_bg, score_fg = "FEE2E2", (127, 29, 29)

    _section_heading(doc, "📊  Match Score", color_rgb=(37, 99, 235))

    score_table = doc.add_table(rows=1, cols=1)
    score_table.style = "Table Grid"
    cell = score_table.cell(0, 0)
    cell.width = Inches(6.5)
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="clear" w:color="auto" w:fill="{score_bg}"/>')
    tcPr.append(shd)

    sp = cell.paragraphs[0]
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp.paragraph_format.space_before = Pt(8)
    sr = sp.add_run(f"{score}%  —  {label}")
    sr.font.name = "Calibri"
    sr.font.size = Pt(24)
    sr.font.bold = True
    sr.font.color.rgb = RGBColor(*score_fg)

    sp2 = cell.add_paragraph()
    sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp2.paragraph_format.space_after = Pt(8)
    sr2 = sp2.add_run(advice)
    sr2.font.name = "Calibri"
    sr2.font.size = Pt(11)
    sr2.font.color.rgb = RGBColor(*score_fg)

    doc.add_paragraph()
    _body_text(doc, f"Reasoning: {gap.get('score_reasoning','')}", italic=True, color_rgb=(100, 116, 139))
    _body_text(doc, f"Education: {gap.get('education_match','')}", italic=True, color_rgb=(100, 116, 139))

    _add_divider(doc, "E2E6F0")
    _section_heading(doc, "✅  Matched Skills", color_rgb=(5, 150, 105))
    _skills_table(doc, gap.get("matched_skills", []), "D1FAE5", (6, 95, 70))

    _section_heading(doc, "⚡  Partial Matches", color_rgb=(217, 119, 6))
    for item in gap.get("partial_skills", []):
        if isinstance(item, dict):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            r1 = p.add_run(f"{item['skill']}: ")
            r1.font.bold = True; r1.font.name = "Calibri"; r1.font.size = Pt(11)
            r1.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
            r2 = p.add_run(f"You have '{item['resume_level']}' — JD needs '{item['required_level']}'")
            r2.font.name = "Calibri"; r2.font.size = Pt(11)
            r2.font.color.rgb = RGBColor(0x37, 0x41, 0x51)

    _section_heading(doc, "❌  Missing Skills", color_rgb=(220, 38, 38))
    _skills_table(doc, gap.get("missing_skills", []), "FEE2E2", (127, 29, 29))
    _add_divider(doc, "E2E6F0")

    _section_heading(doc, "💼  Matching Experience", color_rgb=(5, 150, 105))
    for e in gap.get("matched_experience", []):
        _body_text(doc, f"• {e}")

    _section_heading(doc, "📋  Experience Gaps", color_rgb=(220, 38, 38))
    for e in gap.get("missing_experience", []):
        _body_text(doc, f"• {e}")

    _add_divider(doc, "E2E6F0")

    _section_heading(doc, "💪  Key Strengths", color_rgb=(37, 99, 235))
    for s in gap.get("strengths", []):
        _body_text(doc, f"✨  {s}")

    _section_heading(doc, "🎯  Improvement Suggestions", color_rgb=(124, 58, 237))
    for i, s in enumerate(gap.get("improvement_suggestions", []), 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        rn = p.add_run(f"{i}.  ")
        rn.font.bold = True; rn.font.name = "Calibri"; rn.font.size = Pt(11)
        rn.font.color.rgb = RGBColor(0x7C, 0x3A, 0xED)
        rb = p.add_run(s)
        rb.font.name = "Calibri"; rb.font.size = Pt(11)

    _add_footer(doc, "JobFit AI — Confidential", f"Generated {date_str}")

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


def build_coverletter_docx(cover, applicant, company):
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

    date_str = datetime.now().strftime("%B %d, %Y")
    _doc_header(doc, "Cover Letter", f"Application to {company}", applicant, company, date_str)
    doc.add_paragraph()

    paragraphs = cover.strip().split("\n")
    for para in paragraphs:
        if para.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(para.strip())
            run.font.name = "Calibri"
            run.font.size = Pt(11.5)
            run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x2E)
            p.paragraph_format.line_spacing = Pt(16)

    _add_footer(doc, f"{applicant} — Cover Letter", date_str)
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


def _sanitize(text):
    import re
    text = text.replace('\x00', '')
    text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f]', '', text)
    return text


def build_resume_docx(resume_text, applicant, company):
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

    date_str = datetime.now().strftime("%B %d, %Y")
    _doc_header(doc, applicant, f"Application for: {company}", applicant, company, date_str)
    doc.add_paragraph()

    lines = resume_text.strip().split("\n")
    for line in lines:
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(stripped)
        run.font.name = "Calibri"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1A, 0x1D, 0x2E)

    _add_footer(doc, f"{applicant} — Resume", f"Prepared for {company}  |  {date_str}")
    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ── Analysis runner ───────────────────────────────────────────────────────────
def run_analysis(resume_txt, jd_txt, is_demo=False):
    prog = st.progress(0)
    status = st.empty()

    status.caption("📄 Reading your resume...")
    prog.progress(15)
    time.sleep(0.4)

    status.caption("💼 Parsing job description...")
    prog.progress(30)
    time.sleep(0.4)

    status.caption("🤖 Running AI analysis — this takes ~15 seconds...")
    prog.progress(45)

    try:
        result = analyze(resume_txt, jd_txt)
    except json.JSONDecodeError:
        prog.empty(); status.empty()
        st.error("Failed to parse AI response. Please try again.")
        return
    except Exception as e:
        prog.empty(); status.empty()
        st.error(f"Analysis failed: {e}")
        return

    status.caption("📊 Calculating match score...")
    prog.progress(75)
    time.sleep(0.3)

    status.caption("📝 Generating cover letter...")
    prog.progress(90)
    time.sleep(0.3)

    status.caption("✅ Done!")
    prog.progress(100)
    time.sleep(0.4)

    prog.empty()
    status.empty()

    st.session_state.update({
        "score":          result.get("score", 0),
        "gap_data":       result,
        "cover_letter":   result.get("cover_letter", ""),
        "resume_text":    resume_txt,
        "jd_text":        jd_txt,
        "company_name":   result.get("company_name", "the Company"),
        "applicant_name": result.get("applicant_name", "Applicant"),
        "analysis_done":  True,
        "is_demo":        is_demo,
    })
    st.rerun()


# ── Helper: score badge class ─────────────────────────────────────────────────
def score_badge_class(score):
    if score >= 80:
        return "app-score-green"
    elif score >= 60:
        return "app-score-amber"
    return "app-score-red"


# ── Helper: status badge class ────────────────────────────────────────────────
def status_badge_class(status):
    mapping = {
        "Applied":              "status-applied",
        "Interview Scheduled":  "status-interview",
        "Rejected":             "status-rejected",
        "Offer Received":       "status-offer",
    }
    return mapping.get(status, "status-applied")


# ══════════════════════════════════════════════════════════════════════════════
# TOP BAR (outside tabs — always visible)
# ══════════════════════════════════════════════════════════════════════════════
col_title, col_user = st.columns([4, 1])
with col_user:
    st.markdown(f"""
<div style='text-align:right; padding-top:0.6rem;'>
    <div style='color:#1a1d2e; font-size:1rem; font-weight:600;'>👤 {user_name}</div>
    <div style='color:#94a3b8; font-size:0.78rem;'>{user_email}</div>
</div>""", unsafe_allow_html=True)
    if st.button("Logout", use_container_width=True):
        st.logout()

# Hero
st.markdown("""
<div class="hero">
    <h1>🎯 <span>JobFit AI</span></h1>
    <p>Upload your resume &amp; job description — get a score, gap analysis, and a tailored cover letter instantly.</p>
</div>
""", unsafe_allow_html=True)

# How it works
st.markdown("""
<div class="hiw-wrap">
    <div class="hiw-title">How it works</div>
    <div class="hiw-steps">
        <div class="hiw-step">
            <div class="hiw-icon">📄</div>
            <div class="hiw-step-title">Upload Resume</div>
            <div class="hiw-step-desc">PDF, DOCX or TXT</div>
        </div>
        <div class="hiw-step">
            <div class="hiw-icon">💼</div>
            <div class="hiw-step-title">Add Job Description</div>
            <div class="hiw-step-desc">Upload or paste it</div>
        </div>
        <div class="hiw-step">
            <div class="hiw-icon">🤖</div>
            <div class="hiw-step-title">AI Analyzes</div>
            <div class="hiw-step-desc">Compares both instantly</div>
        </div>
        <div class="hiw-step">
            <div class="hiw-icon">📊</div>
            <div class="hiw-step-title">Get Your Score</div>
            <div class="hiw-step-desc">0–100% match rating</div>
        </div>
        <div class="hiw-step">
            <div class="hiw-icon">📥</div>
            <div class="hiw-step-title">Download Files</div>
            <div class="hiw-step-desc">Report, cover letter &amp; resume</div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)


# ══════════════════════════════════════════════════════════════════════════════
# DEMO BANNER (above tabs — always visible after login)
# ══════════════════════════════════════════════════════════════════════════════
with st.container():
    st.markdown("""
<div style="
    background: linear-gradient(135deg, #eff6ff, #f5f3ff);
    border: 1.5px solid #bfdbfe;
    border-bottom: none;
    border-radius: 12px 12px 0 0;
    padding: 1.2rem 2rem 1rem;
    margin-bottom: 0;
">
    <div style="display:flex; align-items:center; gap:1rem;">
        <div style="font-size:1.8rem; flex-shrink:0;">🧪</div>
        <div>
            <div style="font-weight:700; color:#2563eb; font-size:0.95rem;">New here? Try the demo first</div>
            <div style="color:#64748b; font-size:0.82rem; margin-top:2px;">Load a sample resume &amp; job description with one click — no upload needed.</div>
        </div>
    </div>
</div>
<style>
/* Try Demo button — flush bottom of card */
div[data-testid="stButton"]:has(> button[kind="primary"]) + * { margin-top: 0 !important; }
section[data-testid="stMain"] div[data-testid="stButton"] button[kind="primary"] {
    border-radius: 0 0 12px 12px !important;
    border: 1.5px solid #bfdbfe !important;
    border-top: none !important;
    width: 100%;
    padding-top: 0.65rem !important;
    padding-bottom: 0.65rem !important;
}
</style>
""", unsafe_allow_html=True)
    demo_btn = st.button("▶️ Try Demo", type="primary", use_container_width=True, key="demo_top")

if demo_btn:
    run_analysis(DEMO_RESUME, DEMO_JD, is_demo=True)

# Divider — transition from demo into the main tabs
st.markdown("""
<div style="display:flex; align-items:center; gap:0.6rem; margin:1.2rem 0 0.6rem;">
    <div style="flex:1; height:1px; background:#dde1ef;"></div>
    <div style="font-size:0.75rem; font-weight:700; text-transform:uppercase; letter-spacing:0.1em; color:#94a3b8; white-space:nowrap;">or analyze your own</div>
    <div style="flex:1; height:1px; background:#dde1ef;"></div>
</div>
""", unsafe_allow_html=True)

# ══════════════════════════════════════════════════════════════════════════════
# MAIN NAV TABS
# ══════════════════════════════════════════════════════════════════════════════
tab_analyze, tab_applications, tab_dashboard = st.tabs([
    "⚡ Analyze", "📋 My Applications", "📊 Dashboard"
])


# ══════════════════════════════════════════════════════════════════════════════
# TAB 1 — ANALYZE
# ══════════════════════════════════════════════════════════════════════════════
with tab_analyze:

    # Inputs
    col1, col2 = st.columns(2, gap="large")

    with col1:
        st.markdown('<p style="font-size:0.95rem; font-weight:700; color:#2563eb; margin-bottom:0.3rem;">📄 Your Resume</p>', unsafe_allow_html=True)
        resume_file = st.file_uploader(
            "Upload Resume", type=["pdf", "docx", "txt"],
            key="resume_up", label_visibility="collapsed"
        )
        st.caption("PDF · DOCX · TXT — name extracted automatically")
        resume_text_loaded = ""
        if resume_file:
            resume_text_loaded = extract_text(resume_file)
            if resume_text_loaded:
                st.success("✅ Resume loaded")

    with col2:
        st.markdown('<p style="font-size:0.95rem; font-weight:700; color:#2563eb; margin-bottom:0.3rem;">💼 Job Description</p>', unsafe_allow_html=True)
        jd_upload_tab, jd_paste_tab = st.tabs(["📁 Upload File", "📋 Paste Text"])
        jd_from_file, jd_from_paste = "", ""
        with jd_upload_tab:
            jd_file = st.file_uploader(
                "Upload JD", type=["pdf", "docx", "txt"],
                key="jd_up", label_visibility="collapsed"
            )
            if jd_file:
                jd_from_file = extract_text(jd_file)
                if jd_from_file:
                    st.success("✅ Job description loaded")
        with jd_paste_tab:
            jd_from_paste = st.text_area(
                "Paste JD", height=160,
                placeholder="Paste the full job description here...",
                label_visibility="collapsed"
            )

    jd_final = jd_from_file if jd_from_file.strip() else jd_from_paste

    # CTA row — prominent Analyze button, small Reset on the side
    st.markdown("<div style='margin-top:1rem'></div>", unsafe_allow_html=True)
    cta_col, reset_col, _ = st.columns([2.5, 0.8, 3])
    with cta_col:
        analyze_btn = st.button("⚡ Analyze & Score", use_container_width=True, type="primary")
    with reset_col:
        reset_btn = st.button("↺ Reset", use_container_width=True)

    if reset_btn:
        for k, v in defaults.items():
            st.session_state[k] = v
        st.rerun()

    if analyze_btn:
        if not resume_text_loaded.strip():
            st.error("⚠️ Please upload your resume.")
        elif not jd_final.strip():
            st.error("⚠️ Please upload or paste the job description.")
        else:
            run_analysis(resume_text_loaded, jd_final, is_demo=False)

    # ── Results ───────────────────────────────────────────────────────────────
    if st.session_state.analysis_done and st.session_state.gap_data:
        gap       = st.session_state.gap_data
        score     = st.session_state.score
        company   = st.session_state.company_name
        applicant = st.session_state.applicant_name

        st.markdown("---")

        if st.session_state.get("is_demo"):
            st.info("🧪 **Demo mode** — these are sample results. Upload your own resume and job description above for a real analysis.")

        ic1, ic2, ic3 = st.columns(3)
        with ic1: st.markdown(f'<div class="badge">👤 {applicant}</div>', unsafe_allow_html=True)
        with ic2: st.markdown(f'<div class="badge">🏢 {company}</div>', unsafe_allow_html=True)
        with ic3: st.markdown(f'<div class="badge">📅 {datetime.now().strftime("%b %d, %Y")}</div>', unsafe_allow_html=True)
        st.markdown("")

        res_tab1, res_tab2, res_tab3 = st.tabs(["📊 Score & Overview", "🔍 Gap Analysis", "📝 Cover Letter"])

        with res_tab1:
            card_class, label, advice = score_meta(score)
            left, right = st.columns([1, 2])
            with left:
                st.markdown(f"""
                <div class="score-wrap {card_class}">
                    <div class="score-number">{score}%</div>
                    <div class="score-label">{label}</div>
                    <div class="score-advice">{advice}</div>
                </div>""", unsafe_allow_html=True)
            with right:
                matched = len(gap.get("matched_skills", []))
                partial = len(gap.get("partial_skills", []))
                missing = len(gap.get("missing_skills", []))
                total   = matched + partial + missing or 1
                st.markdown(f"""
                <div class="metrics">
                    <div class="metric"><div class="metric-val" style="color:#059669">{matched}</div><div class="metric-lbl">Matched Skills</div></div>
                    <div class="metric"><div class="metric-val" style="color:#d97706">{partial}</div><div class="metric-lbl">Partial Matches</div></div>
                    <div class="metric"><div class="metric-val" style="color:#dc2626">{missing}</div><div class="metric-lbl">Missing Skills</div></div>
                    <div class="metric"><div class="metric-val" style="color:#2563eb">{round(matched/total*100)}%</div><div class="metric-lbl">Skill Coverage</div></div>
                </div>""", unsafe_allow_html=True)
                st.markdown(f'<div class="card"><h4>Score Reasoning</h4><p>{gap.get("score_reasoning","")}</p></div>', unsafe_allow_html=True)
                st.markdown(f'<div class="card"><h4>Education Fit</h4><p>{gap.get("education_match","")}</p></div>', unsafe_allow_html=True)

            st.markdown("#### 💪 Key Strengths")
            sc = st.columns(3)
            for i, s in enumerate(gap.get("strengths", [])):
                with sc[i % 3]: st.markdown(f'<div class="strength">✨ {s}</div>', unsafe_allow_html=True)

            st.markdown("#### 🎯 Improvement Suggestions")
            for i, sug in enumerate(gap.get("improvement_suggestions", []), 1):
                st.markdown(f'<div class="card"><h4>Suggestion {i}</h4><p>{sug}</p></div>', unsafe_allow_html=True)

        with res_tab2:
            st.markdown("#### ✅ Matched Skills")
            if gap.get("matched_skills"):
                pills = "".join(f'<span class="pill pill-green">{s}</span>' for s in gap["matched_skills"])
                st.markdown(f'<div class="gap-block gap-green"><div class="gap-title">Skills you already have</div>{pills}</div>', unsafe_allow_html=True)
            else:
                st.info("No direct skill matches found.")

            st.markdown("#### ⚡ Partial Matches")
            if gap.get("partial_skills"):
                for item in gap["partial_skills"]:
                    if isinstance(item, dict):
                        st.markdown(f"""
                        <div class="gap-block gap-amber">
                            <div class="gap-title">⚡ {item.get('skill','')}</div>
                            <p class="gap-sub"><b>You have:</b> {item.get('resume_level','')} &nbsp;|&nbsp; <b>JD needs:</b> {item.get('required_level','')}</p>
                        </div>""", unsafe_allow_html=True)
            else:
                st.info("No partial matches found.")

            st.markdown("#### ❌ Missing Skills")
            if gap.get("missing_skills"):
                pills = "".join(f'<span class="pill pill-red">{s}</span>' for s in gap["missing_skills"])
                st.markdown(f'<div class="gap-block gap-red"><div class="gap-title">Skills to develop</div>{pills}</div>', unsafe_allow_html=True)
            else:
                st.success("No critical missing skills!")

            st.markdown("#### 📋 Experience Gaps")
            if gap.get("missing_experience"):
                for exp in gap["missing_experience"]:
                    st.markdown(f'<div class="gap-block gap-red"><p class="gap-sub">• {exp}</p></div>', unsafe_allow_html=True)
            else:
                st.success("Your experience aligns well!")

            st.markdown("#### ✅ Matching Experience")
            if gap.get("matched_experience"):
                for exp in gap["matched_experience"]:
                    st.markdown(f'<div class="gap-block gap-green"><p class="gap-sub">• {exp}</p></div>', unsafe_allow_html=True)

        with res_tab3:
            st.markdown(f"#### 📝 Cover Letter for **{company}**")
            st.markdown(f'<div class="cover-box">{st.session_state.cover_letter}</div>', unsafe_allow_html=True)
            st.markdown("")
            edited_cover = st.text_area(
                "✏️ Edit cover letter before downloading",
                value=st.session_state.cover_letter,
                height=280,
                key="editable_cover",
            )

        # Downloads
        st.markdown("---")
        st.markdown("#### 📥 Download Files")
        st.caption("Three separate files — each ready to use directly.")

        safe_co   = re.sub(r'[^a-zA-Z0-9_-]', '_', company)
        safe_name = re.sub(r'[^a-zA-Z0-9_-]', '_', applicant)
        date_str  = datetime.now().strftime('%Y%m%d')
        cover_use = st.session_state.get("editable_cover", st.session_state.cover_letter)

        if DOCX_OK:
            report_bytes = build_report_docx(score, gap, company)
            cover_bytes  = build_coverletter_docx(cover_use, applicant, company)
            resume_bytes = build_resume_docx(st.session_state.resume_text, applicant, company)

            dl1, dl2, dl3 = st.columns(3)
            with dl1:
                st.markdown('<div class="dl-card"><div class="dl-icon">📊</div><div class="dl-title">Match Report</div><div class="dl-desc">Score · Gap Analysis · Suggestions</div></div>', unsafe_allow_html=True)
                st.download_button("⬇️ Download Report", data=report_bytes,
                    file_name=f"JobFit_Report_{safe_co}_{date_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            with dl2:
                st.markdown('<div class="dl-card"><div class="dl-icon">📝</div><div class="dl-title">Cover Letter</div><div class="dl-desc">Ready to attach to your application</div></div>', unsafe_allow_html=True)
                st.download_button("⬇️ Download Cover Letter", data=cover_bytes,
                    file_name=f"CoverLetter_{safe_co}_{date_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            with dl3:
                st.markdown('<div class="dl-card"><div class="dl-icon">📄</div><div class="dl-title">Resume</div><div class="dl-desc">Tagged with company name</div></div>', unsafe_allow_html=True)
                st.download_button("⬇️ Download Resume", data=resume_bytes,
                    file_name=f"Resume_{safe_name}_{safe_co}_{date_str}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
        else:
            st.warning("Install python-docx to enable downloads: `pip install python-docx`")

        # Save to Tracker
        st.markdown("---")
        st.markdown("#### 💾 Save to Tracker")
        st.caption("Save this application — it will appear in the My Applications tab.")

        tr1, tr2 = st.columns([1, 2])
        with tr1:
            app_status = st.selectbox(
                "Application Status",
                ["Applied", "Interview Scheduled", "Rejected", "Offer Received"],
                key="app_status"
            )
        with tr2:
            app_notes = st.text_input(
                "Notes (optional)",
                placeholder="e.g. Applied via LinkedIn, referral from John...",
                key="app_notes"
            )

        if st.button("💾 Save to Tracker", use_container_width=False):
            try:
                supabase.table("applications").insert({
                    "company":      company,
                    "applicant":    applicant,
                    "score":        score,
                    "status":       app_status,
                    "notes":        app_notes,
                    "date_applied": datetime.now().strftime("%Y-%m-%d"),
                    "user_id":      user_email,
                }).execute()
                st.success(f"✅ Saved! Switch to the **My Applications** tab to see {company}.")
            except Exception as e:
                st.error(f"Failed to save: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 2 — MY APPLICATIONS
# ══════════════════════════════════════════════════════════════════════════════
with tab_applications:
    st.markdown("#### 📋 My Saved Applications")
    st.caption("All your tracked applications — update status, download as CSV, or delete entries.")

    try:
        response = (
            supabase.table("applications")
            .select("*")
            .eq("user_id", user_email)
            .is_("deleted_at", "null")
            .order("date_applied", desc=True)
            .execute()
        )
        data = response.data

        if not data:
            st.info("No applications saved yet. Analyze a job and hit **Save to Tracker** to see it here!")
        else:
            # ── Summary metrics ────────────────────────────────────────────
            total_apps    = len(data)
            avg_score     = round(sum(r["score"] for r in data) / total_apps)
            offers        = sum(1 for r in data if r.get("status") == "Offer Received")
            interviews    = sum(1 for r in data if r.get("status") == "Interview Scheduled")

            m1, m2, m3, m4 = st.columns(4)
            m1.metric("Total Applications", total_apps)
            m2.metric("Avg Match Score", f"{avg_score}%")
            m3.metric("Interviews", interviews)
            m4.metric("Offers", offers)

            st.markdown("")

            # ── Application rows ───────────────────────────────────────────
            for row in data:
                sc       = row.get("score", 0)
                sc_class = score_badge_class(sc)
                st_class = status_badge_class(row.get("status", "Applied"))

                with st.container():
                    left_col, right_col = st.columns([5, 1])

                    with left_col:
                        st.markdown(f"""
<div class="app-row">
    <div class="app-row-company">🏢 {row.get('company','—')}</div>
    <span class="app-score-badge {sc_class}">{sc}%</span>
    <span class="status-badge {st_class}">{row.get('status','—')}</span>
    <div class="app-row-meta">👤 {row.get('applicant','—')} &nbsp;·&nbsp; 📅 {row.get('date_applied','—')}</div>
    {f'<div class="app-row-meta" style="width:100%; margin-top:4px;">📝 {row.get("notes")}</div>' if row.get('notes') else ''}
</div>""", unsafe_allow_html=True)

                    with right_col:
                        st.markdown("<div style='height:12px'></div>", unsafe_allow_html=True)
                        if st.button("🗑️ Delete", key=f"del_{row['id']}", use_container_width=True):
                            try:
                                supabase.table("applications").update({
                                    "deleted_at": datetime.utcnow().isoformat()
                                }).eq("id", row["id"]).execute()
                                st.success(f"Deleted {row.get('company','application')}.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Delete failed: {e}")

            st.markdown("---")

            # ── Update status ──────────────────────────────────────────────
            st.markdown("**✏️ Update Application Status**")
            up1, up2, up3 = st.columns([2, 2, 1])
            with up1:
                companies = [row["company"] for row in data]
                selected_company = st.selectbox("Select Company", companies, key="update_company")
            with up2:
                new_status = st.selectbox(
                    "New Status",
                    ["Applied", "Interview Scheduled", "Rejected", "Offer Received"],
                    key="new_status"
                )
            with up3:
                st.markdown("<br>", unsafe_allow_html=True)
                if st.button("Update", use_container_width=True):
                    try:
                        selected_id = next(row["id"] for row in data if row["company"] == selected_company)
                        supabase.table("applications").update({"status": new_status}).eq("id", selected_id).execute()
                        st.success(f"✅ Updated {selected_company} to '{new_status}'")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Update failed: {e}")

            st.markdown("")

            # ── CSV download ───────────────────────────────────────────────
            df = pd.DataFrame(data)
            df = df[["date_applied", "company", "applicant", "score", "status", "notes"]]
            df.columns = ["Date", "Company", "Applicant", "Score (%)", "Status", "Notes"]
            csv_bytes = df.to_csv(index=False).encode("utf-8")
            st.download_button(
                "⬇️ Download All as CSV",
                data=csv_bytes,
                file_name=f"JobFit_Tracker_{datetime.now().strftime('%Y%m%d')}.csv",
                mime="text/csv",
            )

    except Exception as e:
        st.error(f"Could not load applications: {e}")


# ══════════════════════════════════════════════════════════════════════════════
# TAB 3 — DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
with tab_dashboard:
    import plotly.express as px
    import plotly.graph_objects as go

    st.markdown("#### 📊 Dashboard")
    st.caption("Your personal job application analytics — updated every time you save an application.")

    # ── Fetch data ─────────────────────────────────────────────────────────────
    try:
        dash_resp = (
            supabase.table("applications")
            .select("*")
            .eq("user_id", user_email)
            .is_("deleted_at", "null")
            .order("date_applied", desc=False)
            .execute()
        )
        dash_data = dash_resp.data
    except Exception as e:
        st.error(f"Could not load dashboard data: {e}")
        dash_data = []

    if not dash_data:
        st.markdown("""
<div style="
    background: #ffffff; border: 1px solid #dde1ef; border-radius: 14px;
    padding: 3rem 2rem; text-align: center; margin-top: 1rem;
">
    <div style="font-size:3rem; margin-bottom:1rem;">📭</div>
    <div style="font-size:1.1rem; font-weight:700; color:#1a1d2e; margin-bottom:0.4rem;">No data yet</div>
    <div style="color:#64748b; font-size:0.9rem;">Analyze a job and save it to the tracker — your dashboard will come alive automatically.</div>
</div>
""", unsafe_allow_html=True)

    else:
        df = pd.DataFrame(dash_data)
        df = df.copy()
        df.loc[:, "score"]        = pd.to_numeric(df["score"], errors="coerce").fillna(0)
        df.loc[:, "date_applied"] = pd.to_datetime(df["date_applied"], errors="coerce")

        total_apps = len(df)
        avg_score  = round(df["score"].mean())
        best_score = int(df["score"].max())
        interviews = int((df["status"] == "Interview Scheduled").sum())
        offers     = int((df["status"] == "Offer Received").sum())
        rejections = int((df["status"] == "Rejected").sum())

        # ── Stat cards ─────────────────────────────────────────────────────────
        s1, s2, s3, s4, s5 = st.columns(5)
        s1.metric("Total Apps",    total_apps)
        s2.metric("Avg Score",     f"{avg_score}%")
        s3.metric("Best Score",    f"{best_score}%")
        s4.metric("Interviews",    interviews)
        s5.metric("Offers",        offers)

        st.markdown("")

        COLORS = {
            "Applied":             "#3b82f6",
            "Interview Scheduled": "#7c3aed",
            "Offer Received":      "#059669",
            "Rejected":            "#dc2626",
        }

        # ── Row 2: Bar chart + Donut ────────────────────────────────────────────
        chart_left, chart_right = st.columns([3, 2])

        with chart_left:
            st.markdown('<p style="font-weight:700; font-size:0.9rem; color:#1a1d2e; margin-bottom:0.3rem;">🏢 Match Score by Company</p>', unsafe_allow_html=True)
            df_bar = df.sort_values("score", ascending=True).copy()
            bar_colors = df_bar["score"].apply(
                lambda s: "#059669" if s >= 80 else ("#d97706" if s >= 60 else "#dc2626")
            )
            fig_bar = go.Figure(go.Bar(
                x=df_bar["score"],
                y=df_bar["company"],
                orientation="h",
                marker_color=bar_colors,
                text=df_bar["score"].apply(lambda s: f"{s}%"),
                textposition="outside",
                hovertemplate="<b>%{y}</b><br>Score: %{x}%<extra></extra>",
            ))
            fig_bar.update_layout(
                margin=dict(l=0, r=40, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                xaxis=dict(
                    showgrid=True, gridcolor="#f1f5f9",
                    range=[0, 110], showticklabels=False,
                    zeroline=False,
                ),
                yaxis=dict(showgrid=False, tickfont=dict(size=12)),
                height=max(200, len(df_bar) * 52),
                showlegend=False,
            )
            st.plotly_chart(fig_bar, width="stretch")

        with chart_right:
            st.markdown('<p style="font-weight:700; font-size:0.9rem; color:#1a1d2e; margin-bottom:0.3rem;">📋 Application Status</p>', unsafe_allow_html=True)
            status_counts = df["status"].value_counts().reset_index()
            status_counts.columns = ["status", "count"]
            donut_colors = [COLORS.get(s, "#94a3b8") for s in status_counts["status"]]
            fig_donut = go.Figure(go.Pie(
                labels=status_counts["status"],
                values=status_counts["count"],
                hole=0.55,
                marker_colors=donut_colors,
                textinfo="label+percent",
                hovertemplate="<b>%{label}</b><br>%{value} application(s)<extra></extra>",
            ))
            fig_donut.update_layout(
                margin=dict(l=0, r=0, t=10, b=10),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                showlegend=False,
                height=max(200, len(df_bar) * 52),
            )
            st.plotly_chart(fig_donut, width="stretch")

        # ── Row 3: Score over time ──────────────────────────────────────────────
        st.markdown('<p style="font-weight:700; font-size:0.9rem; color:#1a1d2e; margin-bottom:0.3rem;">📈 Score Over Time</p>', unsafe_allow_html=True)

        df_time = df.dropna(subset=["date_applied"]).sort_values("date_applied")

        if len(df_time) < 2:
            st.info("Apply to at least 2 jobs to see your score trend over time.")
        else:
            status_color_list = [COLORS.get(s, "#3b82f6") for s in df_time["status"]]
            fig_line = go.Figure()

            # Shaded reference bands
            fig_line.add_hrect(y0=80, y1=100, fillcolor="#d1fae5", opacity=0.25, line_width=0)
            fig_line.add_hrect(y0=60, y1=80,  fillcolor="#fef3c7", opacity=0.25, line_width=0)
            fig_line.add_hrect(y0=0,  y1=60,  fillcolor="#fee2e2", opacity=0.2,  line_width=0)

            # Line
            fig_line.add_trace(go.Scatter(
                x=df_time["date_applied"],
                y=df_time["score"],
                mode="lines+markers+text",
                line=dict(color="#2563eb", width=2.5),
                marker=dict(
                    size=10,
                    color=status_color_list,
                    line=dict(color="white", width=2),
                ),
                text=df_time["company"],
                textposition="top center",
                textfont=dict(size=11),
                hovertemplate="<b>%{text}</b><br>Score: %{y}%<br>Date: %{x|%b %d, %Y}<extra></extra>",
            ))

            # Reference lines
            fig_line.add_hline(y=80, line_dash="dot", line_color="#059669",
                               annotation_text="Strong match", annotation_position="right",
                               annotation_font_size=11)
            fig_line.add_hline(y=60, line_dash="dot", line_color="#d97706",
                               annotation_text="Decent match", annotation_position="right",
                               annotation_font_size=11)

            fig_line.update_layout(
                margin=dict(l=0, r=80, t=20, b=20),
                paper_bgcolor="rgba(0,0,0,0)",
                plot_bgcolor="rgba(0,0,0,0)",
                xaxis=dict(showgrid=False, zeroline=False, tickformat="%b %d"),
                yaxis=dict(showgrid=True, gridcolor="#f1f5f9", range=[0, 105],
                           ticksuffix="%", zeroline=False),
                height=320,
                showlegend=False,
            )
            st.plotly_chart(fig_line, width="stretch")

            # Legend for dot colors
            st.markdown("""
<div style="display:flex; gap:1.2rem; font-size:0.78rem; color:#64748b; margin-top:-0.5rem; flex-wrap:wrap;">
    <span><span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#3b82f6;margin-right:4px;"></span>Applied</span>
    <span><span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#7c3aed;margin-right:4px;"></span>Interview</span>
    <span><span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#059669;margin-right:4px;"></span>Offer</span>
    <span><span style="display:inline-block;width:10px;height:10px;border-radius:50%;background:#dc2626;margin-right:4px;"></span>Rejected</span>
</div>
""", unsafe_allow_html=True)

        # ── Row 4: Top missing skills ───────────────────────────────────────────
        st.markdown("")
        st.markdown('<p style="font-weight:700; font-size:0.9rem; color:#1a1d2e; margin-bottom:0.3rem;">🔍 Quick Stats</p>', unsafe_allow_html=True)

        q1, q2, q3 = st.columns(3)
        with q1:
            best_row = df.loc[df["score"].idxmax()]
            st.markdown(f"""
<div class="card">
    <h4>🏆 Best Match</h4>
    <p style="font-size:1.1rem; font-weight:700; color:#059669;">{int(best_row['score'])}% — {best_row['company']}</p>
    <p style="font-size:0.8rem; margin-top:4px;">{best_row['date_applied'].strftime('%b %d, %Y') if pd.notna(best_row['date_applied']) else '—'}</p>
</div>""", unsafe_allow_html=True)

        with q2:
            worst_row = df.loc[df["score"].idxmin()]
            st.markdown(f"""
<div class="card">
    <h4>📉 Lowest Match</h4>
    <p style="font-size:1.1rem; font-weight:700; color:#dc2626;">{int(worst_row['score'])}% — {worst_row['company']}</p>
    <p style="font-size:0.8rem; margin-top:4px;">{worst_row['date_applied'].strftime('%b %d, %Y') if pd.notna(worst_row['date_applied']) else '—'}</p>
</div>""", unsafe_allow_html=True)

        with q3:
            above_80 = int((df["score"] >= 80).sum())
            pct = round(above_80 / total_apps * 100)
            st.markdown(f"""
<div class="card">
    <h4>🎯 Strong Matches</h4>
    <p style="font-size:1.1rem; font-weight:700; color:#2563eb;">{above_80} of {total_apps} apps ({pct}%)</p>
    <p style="font-size:0.8rem; margin-top:4px;">scored 80% or above</p>
</div>""", unsafe_allow_html=True)
